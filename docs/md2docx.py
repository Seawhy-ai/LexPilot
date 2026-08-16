# -*- coding: utf-8 -*-
"""把 docs/ 下的 .md 转成 Word(.docx)。
用法:python md2docx.py [文件1.md 文件2.md ...]   (缺省转换同目录全部 .md)
"""
import re, os, sys, glob
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image as _PIL

_HERE = os.path.dirname(os.path.abspath(__file__))
ARGS = sys.argv[1:]
SRCS = ARGS if ARGS else [f for f in glob.glob(os.path.join(_HERE, '*.md')) if not f.endswith('-formatted.md')]

BOLD_RE = re.compile(r'\*\*(.+?)\*\*')
CODE_RE = re.compile(r'`(.+?)`')


def add_runs(par, text):
    """把 **加粗** 和 `代码` 转成带格式的 run。"""
    tokens = BOLD_RE.split(text)
    for i, tok in enumerate(tokens):
        if tok == '':
            continue
        if i % 2 == 1:
            par.add_run(tok).bold = True
        else:
            parts = CODE_RE.split(tok)
            for j, sub in enumerate(parts):
                if sub == '':
                    continue
                r = par.add_run(sub)
                if j % 2 == 1:
                    r.font.name = 'Consolas'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(10)


def convert(src):
    doc = Document()

    # 基础字体:西文 Calibri / 中文宋体
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 标题样式:黑色加粗
    for name, size in [('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 12)]:
        st = doc.styles[name]
        st.font.name = 'Calibri'
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    with open(src, encoding='utf-8') as f:
        lines = f.read().split('\n')

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # 代码块
        if line.startswith('```'):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code.append(lines[i].rstrip())
                i += 1
            par = doc.add_paragraph()
            r = par.add_run('\n'.join(code))
            r.font.name = 'Consolas'
            r.font.size = Pt(9)
            par.paragraph_format.left_indent = Pt(10)
            i += 1
            continue
        # 空行 / 水平线
        if line.strip() == '':
            i += 1
            continue
        if re.match(r'^\s*-{3,}\s*$', line):
            i += 1
            continue
        # 图片
        if line.startswith('!['):
            m = re.match(r'!\[(.*?)\]\((.+?)\)', line)
            if m:
                alt, rel = m.group(1), m.group(2)
                path = os.path.join(os.path.dirname(src), rel)
                if os.path.exists(path):
                    iw, ih = _PIL.open(path).size
                    width = Cm(12) if ih > iw else Cm(15)
                    doc.add_picture(path, width=width)
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if alt:
                        cap = doc.add_paragraph(alt)
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in cap.runs:
                            r.font.size = Pt(9)
                            r.italic = True
                            r.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
                else:
                    doc.add_paragraph('[图片缺失: %s]' % rel)
            i += 1
            continue
        # 表格
        if line.strip().startswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append(lines[i].strip())
                i += 1
            data = []
            for rline in rows:
                cells = [c.strip() for c in rline.strip('|').split('|')]
                if all(re.fullmatch(r'[\s:\-]+', c) for c in cells):
                    continue  # 表头分隔行
                data.append(cells)
            if data:
                table = doc.add_table(rows=len(data), cols=max(len(r) for r in data))
                table.style = 'Table Grid'
                for ri, cells in enumerate(data):
                    for ci, c in enumerate(cells):
                        cell = table.cell(ri, ci)
                        cp = cell.paragraphs[0]
                        add_runs(cp, c)
                        if ri == 0:
                            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in cp.runs:
                                run.bold = True
            continue
        # 标题
        if line.startswith('# '):
            par = doc.add_heading(line[2:].strip(), 0)
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), 3)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), 2)
            i += 1
            continue
        # 引用
        if line.startswith('> '):
            par = doc.add_paragraph()
            add_runs(par, line[2:].strip())
            par.paragraph_format.left_indent = Pt(12)
            for r in par.runs:
                r.italic = True
            i += 1
            continue
        # 无序列表
        if line.startswith('- '):
            par = doc.add_paragraph(style='List Bullet')
            add_runs(par, line[2:].strip())
            i += 1
            continue
        # 有序列表
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            par = doc.add_paragraph()
            add_runs(par, m.group(1) + '. ' + m.group(2))
            par.paragraph_format.left_indent = Pt(18)
            i += 1
            continue
        # 普通段落
        par = doc.add_paragraph()
        add_runs(par, line)
        i += 1

    dst = src[:-3] + '.docx'
    doc.save(dst)
    print('saved:', dst)


for s in SRCS:
    convert(s)
